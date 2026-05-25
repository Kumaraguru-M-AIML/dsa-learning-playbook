import docx
import os

def extract_project_data(file_path):
    doc = docx.Document(file_path)
    full_text = []
    for para in doc.paragraphs:
        full_text.append(para.text)
    
    # Also attempt to get table data
    for table in doc.tables:
        for row in table.rows:
            row_text = [cell.text for cell in row.cells]
            full_text.append(" | ".join(row_text))
            
    return "\n".join(full_text)

if __name__ == "__main__":
    report_path = r"e:\Wk 7 - Carrier mastery and R&D\Career_Mastery_Hub\04_Experience_and_Projects\College_Projects\Decentralized_Social_Media_Project\project_report_J7.docx"
    output_path = "project_context_full.txt"
    
    try:
        content = extract_project_data(report_path)
        with open(output_path, "w", encoding="utf-8") as f:
            f.write(content)
        print(f"Extracted {len(content)} characters to {output_path}")
    except Exception as e:
        print(f"Error: {e}")
